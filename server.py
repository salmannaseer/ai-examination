"""
Complete FastAPI Server for Sindh Ombudsman AI Prototype.
Handles scanned PDFs, digital PDFs, DOCX, TXT, and JPG/PNG documents.
Renders scanned PDF pages and images to Base64 multimodal vision parts for GPT-5.4-mini.
"""

import io
import os
import re
import json
import base64
import traceback
from typing import List, Optional, Tuple

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
import httpx
import uvicorn

# Parsing libraries
import fitz  # PyMuPDF
import docx  # python-docx
from PIL import Image

app = FastAPI(
    title="Sindh Ombudsman AI Server",
    version="2.2.0",
    description="Multimodal Base64 Vision & Document Server for Sindh Ombudsman."
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

HTML_FILE_PATH = os.path.join(os.path.dirname(__file__), "Sindh Ombudsman AI Prototype.html")
ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "jpg", "jpeg", "png"}


def validate_file_extension(filename: str) -> str:
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File '{filename}' has an unsupported extension (.{ext}). Only PDF, Word (.docx/.doc), Text (.txt), and Images (.jpg/.png) are allowed."
        )
    return ext


def get_mime_type(ext: str) -> str:
    mime_map = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "txt": "text/plain",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png"
    }
    return mime_map.get(ext, "application/octet-stream")


def process_pdf_document(file_bytes: bytes, max_pages: int = 10) -> Tuple[str, List[dict], int]:
    """
    Process PDF documents (both digital text and scanned photo/image pages).
    Extracts text if available AND renders pages as Base64 images for GPT-5.4-mini Vision.
    """
    text_parts = []
    page_images = []
    total_pages = 0

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        pages_to_render = min(total_pages, max_pages)

        for p_num in range(pages_to_render):
            page = doc[p_num]
            # 1. Digital text extraction (if embedded)
            page_text = page.get_text("text").strip()
            if page_text:
                text_parts.append(f"--- Page {p_num + 1} Text ---\n{page_text}")

            # 2. Render high-resolution page image for visual AI analysis (OCR/Scanned letters)
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("jpeg")
            b64_page = base64.b64encode(img_bytes).decode("utf-8")
            page_images.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{b64_page}",
                    "detail": "high"
                }
            })

        doc.close()
        combined_text = "\n\n".join(text_parts).strip()
        return combined_text, page_images, total_pages
    except Exception as e:
        return f"[PDF parsing notice: {str(e)}]", [], 0


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text and tables from Word (.docx)."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                parts.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(table_rows))
        return "\n\n".join(parts).strip()
    except Exception as e:
        return f"[Word document parsing notice: {str(e)}]"


def extract_plain_text(file_bytes: bytes) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def parse_model_params(param_str: str) -> dict:
    out = {}
    if not param_str:
        return out
    for p in re.split(r'[,;\n]', param_str):
        parts = p.split('=')
        if len(parts) == 2:
            k = parts[0].strip()
            try:
                out[k] = float(parts[1].strip())
            except ValueError:
                pass
    return out


async def call_llm_multimodal(
    endpoint: str,
    api_key: str,
    model: str,
    system_prompt: str,
    user_text: str,
    multimodal_parts: List[dict] = [],
    max_tokens: int = 4000,
    org_id: str = "",
    model_params_str: str = "temperature=0.2",
    timeout_sec: int = 60
) -> str:
    if not api_key:
        raise HTTPException(status_code=400, detail="No API key provided. Please configure your API key in AI Settings.")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    if org_id.strip():
        headers["OpenAI-Organization"] = org_id.strip()

    extra_params = parse_model_params(model_params_str)
    cap = max(max_tokens, 512)

    # Combine text prompt with visual/scanned Base64 page images
    if multimodal_parts:
        user_content = [{"type": "text", "text": user_text}] + multimodal_parts
    else:
        user_content = user_text

    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ],
        "max_tokens": cap,
        **extra_params
    }

    async with httpx.AsyncClient(timeout=float(timeout_sec)) as client:
        try:
            resp = await client.post(endpoint, headers=headers, json=body)

            # Retry with max_completion_tokens for GPT-5 / newer reasoning models
            if resp.status_code != 200 and ("max_completion_tokens" in resp.text or "max_tokens" in resp.text):
                body.pop("max_tokens", None)
                body["max_completion_tokens"] = cap
                resp = await client.post(endpoint, headers=headers, json=body)

            # Retry without temperature if model rejects temperature
            if resp.status_code != 200 and "temperature" in resp.text:
                body.pop("temperature", None)
                resp = await client.post(endpoint, headers=headers, json=body)

            if resp.status_code != 200:
                err_detail = resp.text[:300]
                try:
                    j = resp.json()
                    err_detail = j.get("error", {}).get("message", err_detail)
                except Exception:
                    pass
                raise HTTPException(status_code=resp.status_code, detail=f"LLM Provider Error (HTTP {resp.status_code}): {err_detail}")

            data = resp.json()
            choices = data.get("choices", [])
            if not choices:
                raise HTTPException(status_code=500, detail="LLM provider returned no choices.")

            msg = choices[0].get("message", {})
            content = msg.get("content", "")
            if isinstance(content, list):
                content = "".join([c if isinstance(c, str) else c.get("text", "") for c in content])
            if not content or not str(content).strip():
                content = msg.get("reasoning_content") or choices[0].get("text", "")

            if not content or not str(content).strip():
                raise HTTPException(status_code=500, detail="LLM accepted request but returned empty text.")

            return content.strip()

        except httpx.TimeoutException:
            raise HTTPException(status_code=504, detail=f"Request to AI provider timed out after {timeout_sec}s.")
        except httpx.ConnectError:
            raise HTTPException(status_code=502, detail=f"Could not connect to AI endpoint: {endpoint}")



SETTINGS_FILE_PATH = os.path.join(os.path.dirname(__file__), "settings.json")


@app.get("/api/settings")
async def get_settings():
    if os.path.exists(SETTINGS_FILE_PATH):
        try:
            with open(SETTINGS_FILE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


@app.post("/api/settings")
async def save_settings(settings: dict):
    try:
        # Strip API key from server-side persistence - API key lives in browser only
        clean_settings = {k: v for k, v in settings.items() if k != "apiKey"}
        with open(SETTINGS_FILE_PATH, "w", encoding="utf-8") as f:
            json.dump(clean_settings, f, indent=2)
        return {"success": True}
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})

# =====================================================================
# API Endpoints
# =====================================================================

@app.get("/")
async def serve_prototype():
    if os.path.exists(HTML_FILE_PATH):
        return FileResponse(HTML_FILE_PATH, media_type="text/html")
    return HTMLResponse("<h3>Sindh Ombudsman AI Prototype HTML file not found.</h3>", status_code=404)


@app.get("/api/health")
async def health_check():
    return {
        "status": "online",
        "service": "Sindh Ombudsman AI Server",
        "version": "2.2.0",
        "allowed_formats": ["pdf", "docx", "doc", "txt", "jpg", "jpeg", "png"],
        "scanned_pdf_rendering": True,
        "base64_multimodal_vision": True
    }


@app.post("/api/ai/test-connection")
async def api_test_connection(
    endpoint: str = Form(...),
    apiKey: str = Form(...),
    model: str = Form("gpt-5.4-mini"),
    orgId: str = Form(""),
    timeout: str = Form("30"),
    modelParams: str = Form("temperature=0.2")
):
    try:
        reply = await call_llm_multimodal(
            endpoint=endpoint,
            api_key=apiKey,
            model=model,
            system_prompt="Reply with the single word OK.",
            user_text="Connection test.",
            max_tokens=512,
            org_id=orgId,
            model_params_str=modelParams,
            timeout_sec=int(timeout) if timeout.isdigit() else 30
        )
        return {"success": True, "model": model, "reply": reply}
    except HTTPException as he:
        return JSONResponse(status_code=he.status_code, content={"success": False, "error": he.detail})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@app.post("/api/ai/examine")
async def api_examine(
    complaintType: str = Form(""),
    regionalOffice: str = Form(""),
    complainantName: str = Form(""),
    province: str = Form("Sindh"),
    district: str = Form(""),
    taluka: str = Form(""),
    complaintSubject: str = Form(""),
    complaintDetails: str = Form(""),
    globalInstructions: str = Form(""),
    instructions: str = Form(""),
    endpoint: str = Form("https://api.openai.com/v1/chat/completions"),
    apiKey: str = Form(""),
    model: str = Form("gpt-5.4-mini"),
    orgId: str = Form(""),
    timeout: str = Form("60"),
    maxTokens: str = Form("2500"),
    modelParams: str = Form("temperature=0.2"),
    files: List[UploadFile] = File([])
):
    file_records = []
    multimodal_parts = []

    for i, f in enumerate(files):
        ext = validate_file_extension(f.filename)
        mime = get_mime_type(ext)
        content_bytes = await f.read()

        # 1. PDF (Scanned or Digital): Extracts digital text AND renders visual page images
        if ext == "pdf":
            p_text, p_images, num_pages = process_pdf_document(content_bytes, max_pages=6)
            multimodal_parts.extend(p_images)
            text_note = f"\nExtracted Text:\n{p_text[:6000]}" if p_text else "\nNote: Scanned document attached as high-resolution visual pages for AI inspection."
            file_records.append(f"File {i + 1}: {f.filename} ({num_pages} pages, {len(content_bytes)} bytes){text_note}")

        # 2. Images (JPG / PNG): Base64 visual part
        elif ext in ["jpg", "jpeg", "png"]:
            b64_data = base64.b64encode(content_bytes).decode("utf-8")
            multimodal_parts.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime};base64,{b64_data}",
                    "detail": "high"
                }
            })
            file_records.append(f"File {i + 1}: {f.filename} (Attached Image Document, {len(content_bytes)} bytes)")

        # 3. Word (.docx / .doc): Extract text & tables
        elif ext in ["docx", "doc"]:
            text_extracted = extract_docx_text(content_bytes)
            file_records.append(f"File {i + 1}: {f.filename} (Word Document, {len(content_bytes)} bytes)\nExtracted Content:\n{text_extracted[:6000]}")

        # 4. Text (.txt): Decode UTF-8
        elif ext == "txt":
            text_extracted = extract_plain_text(content_bytes)
            file_records.append(f"File {i + 1}: {f.filename} (Text Document, {len(content_bytes)} bytes)\nExtracted Content:\n{text_extracted[:6000]}")

    files_text = "\n\n".join(file_records) if file_records else "None supplied."

    complaint_block = f"""Type of Complaint: {complaintType or '(not supplied)'}
Regional Office: {regionalOffice or '(not supplied)'}
Complainant Name: {complainantName or '(not supplied)'}
Province: {province or '(not supplied)'}
District: {district or '(not supplied)'}
Taluka: {taluka or '(not supplied)'}
Complaint Subject: {complaintSubject or '(not supplied)'}
Complaint Details:
{complaintDetails or '(not supplied)'}

Attached Files & Documents (Includes Visual Scanned Pages for Vision Analysis):
{files_text}"""

    system_prompt = f"""{globalInstructions}

{instructions}

CRITICAL FORMAT REQUIREMENT:
Carefully inspect the complaint narrative and all attached scanned documents, visual pages, and notices.
You MUST reply with a single valid JSON object only. No conversational text outside JSON.
Keys:
- "recommendation": exactly one of ["May be Admitted", "May be Not Admitted", "May be Forwarded", "May be Deferred"]
- "remarks": detailed examination remarks in plain text referencing the complaint facts and attached evidence
- "notification": notification message for the complainant (max 200 characters)
- "officer_message": message for the Investigation Officer (max 200 characters)"""

    user_text = f"COMPLAINT RECORD\n\n{complaint_block}"

    try:
        raw_response = await call_llm_multimodal(
            endpoint=endpoint,
            api_key=apiKey,
            model=model,
            system_prompt=system_prompt,
            user_text=user_text,
            multimodal_parts=multimodal_parts,
            max_tokens=int(maxTokens) if maxTokens.isdigit() else 2500,
            org_id=orgId,
            model_params_str=modelParams,
            timeout_sec=int(timeout) if timeout.isdigit() else 60
        )

        cleaned = re.sub(r'```json\s*', '', raw_response, flags=re.IGNORECASE)
        cleaned = re.sub(r'```\s*', '', cleaned).strip()
        m = re.search(r'\{[\s\S]*\}', cleaned)
        data = None
        if m:
            try:
                data = json.loads(m.group(0))
            except Exception:
                pass
        if not data:
            try:
                data = json.loads(cleaned)
            except Exception:
                pass

        if not data or not isinstance(data, dict):
            rec = "May be Deferred"
            if re.search(r'May\s+be\s+Admitted', raw_response, re.I):
                rec = "May be Admitted"
            elif re.search(r'May\s+be\s+Not\s+Admitted', raw_response, re.I):
                rec = "May be Not Admitted"
            elif re.search(r'May\s+be\s+Forwarded', raw_response, re.I):
                rec = "May be Forwarded"

            data = {
                "recommendation": rec,
                "remarks": raw_response.strip(),
                "notification": (raw_response[:190] + "...").strip(),
                "officer_message": "Please review the attached scanned documents and case record."
            }

        return {
            "success": True,
            "recommendation": data.get("recommendation", "May be Deferred"),
            "remarks": data.get("remarks", "").strip(),
            "notification": data.get("notification", "").strip(),
            "officer_message": data.get("officer_message", "").strip()
        }

    except HTTPException as he:
        return JSONResponse(status_code=he.status_code, content={"success": False, "error": he.detail})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@app.post("/api/ai/findings")
async def api_findings(
    complaintBlockText: str = Form(""),
    recommendation: str = Form(""),
    remarks: str = Form(""),
    investigationDetails: str = Form(""),
    fileDescriptionsJson: str = Form("[]"),
    globalInstructions: str = Form(""),
    instructions: str = Form(""),
    endpoint: str = Form("https://api.openai.com/v1/chat/completions"),
    apiKey: str = Form(""),
    model: str = Form("gpt-5.4-mini"),
    orgId: str = Form(""),
    timeout: str = Form("60"),
    maxTokens: str = Form("6000"),
    modelParams: str = Form("temperature=0.2"),
    files: List[UploadFile] = File([])
):
    try:
        desc_list = json.loads(fileDescriptionsJson)
    except Exception:
        desc_list = []

    file_records = []
    multimodal_parts = []

    for i, f in enumerate(files):
        ext = validate_file_extension(f.filename)
        mime = get_mime_type(ext)
        content_bytes = await f.read()
        desc = desc_list[i] if i < len(desc_list) else ""
        desc_text = f"\nDescription: {desc}" if desc else ""

        if ext == "pdf":
            p_text, p_images, num_pages = process_pdf_document(content_bytes, max_pages=6)
            multimodal_parts.extend(p_images)
            text_note = f"\nExtracted Text:\n{p_text[:6000]}" if p_text else "\nNote: Scanned record attached as visual pages for vision inspection."
            file_records.append(f"File {i + 1}: {f.filename} ({num_pages} pages){desc_text}{text_note}")

        elif ext in ["jpg", "jpeg", "png"]:
            b64_data = base64.b64encode(content_bytes).decode("utf-8")
            multimodal_parts.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime};base64,{b64_data}",
                    "detail": "high"
                }
            })
            file_records.append(f"File {i + 1}: {f.filename} (Image Record){desc_text}")

        elif ext in ["docx", "doc"]:
            text_extracted = extract_docx_text(content_bytes)
            file_records.append(f"File {i + 1}: {f.filename} (Word){desc_text}\nExtracted Text:\n{text_extracted[:6000]}")

        elif ext == "txt":
            text_extracted = extract_plain_text(content_bytes)
            file_records.append(f"File {i + 1}: {f.filename} (Text){desc_text}\nExtracted Text:\n{text_extracted[:6000]}")

    inv_files_text = "\n\n".join(file_records) if file_records else "None supplied."

    investigation_block = f"""Investigation Details:
{investigationDetails or '(not supplied)'}

Investigation Files and Descriptions:
{inv_files_text}"""

    system_prompt = f"""{globalInstructions}

{instructions}

Return an HTML fragment only: each heading as <h4>, body text as <p>, lists as <ul><li>.
Headings must be exactly: Background, Admissibility, Investigation, Findings, Conclusion.
No markdown code fences, no <html> or <body> tags."""

    user_text = f"""COMPLAINT RECORD

{complaintBlockText}

EXAMINATION
Recommendation: {recommendation or '(not generated)'}
Examination Remarks:
{remarks or '(not generated)'}

INVESTIGATION
{investigation_block}"""

    try:
        raw_response = await call_llm_multimodal(
            endpoint=endpoint,
            api_key=apiKey,
            model=model,
            system_prompt=system_prompt,
            user_text=user_text,
            multimodal_parts=multimodal_parts,
            max_tokens=int(maxTokens) if maxTokens.isdigit() else 6000,
            org_id=orgId,
            model_params_str=modelParams,
            timeout_sec=int(timeout) if timeout.isdigit() else 60
        )
        html = re.sub(r'```html?\s*', '', raw_response, flags=re.IGNORECASE)
        html = re.sub(r'```\s*', '', html).strip()
        return {"success": True, "findingsHtml": html}

    except HTTPException as he:
        return JSONResponse(status_code=he.status_code, content={"success": False, "error": he.detail})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@app.post("/api/ai/draft")
async def api_draft(
    complaintBlockText: str = Form(""),
    investigationBlockText: str = Form(""),
    findingsText: str = Form(""),
    decisionDirections: str = Form(""),
    globalInstructions: str = Form(""),
    instructions: str = Form(""),
    endpoint: str = Form("https://api.openai.com/v1/chat/completions"),
    apiKey: str = Form(""),
    model: str = Form("gpt-5.4-mini"),
    orgId: str = Form(""),
    timeout: str = Form("60"),
    maxTokens: str = Form("8000"),
    modelParams: str = Form("temperature=0.2")
):
    system_prompt = f"""{globalInstructions}

{instructions}

Return an HTML fragment only: each heading as <h4>, body text as <p>, proceedings as <ul><li>.
Headings must be exactly: "1. Complaint", "2. Proceedings", "3. Findings", "4. Decision".
No markdown code fences."""

    user_text = f"""COMPLAINT CONTEXT

{complaintBlockText}

INVESTIGATION CONTEXT
{investigationBlockText}

FINDINGS CONTEXT (latest user-reviewed version)
{findingsText or '(not generated)'}

DECISION CONTEXT — authoritative decision directions entered by the officer
{decisionDirections}"""

    try:
        raw_response = await call_llm_multimodal(
            endpoint=endpoint,
            api_key=apiKey,
            model=model,
            system_prompt=system_prompt,
            user_text=user_text,
            max_tokens=int(maxTokens) if maxTokens.isdigit() else 8000,
            org_id=orgId,
            model_params_str=modelParams,
            timeout_sec=int(timeout) if timeout.isdigit() else 60
        )
        html = re.sub(r'```html?\s*', '', raw_response, flags=re.IGNORECASE)
        html = re.sub(r'```\s*', '', html).strip()
        return {"success": True, "draftHtml": html}

    except HTTPException as he:
        return JSONResponse(status_code=he.status_code, content={"success": False, "error": he.detail})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@app.post("/api/ai/regen-piece")
async def api_regen_piece(
    which: str = Form(...),
    complaintBlockText: str = Form(""),
    recommendation: str = Form(""),
    remarks: str = Form(""),
    globalInstructions: str = Form(""),
    endpoint: str = Form("https://api.openai.com/v1/chat/completions"),
    apiKey: str = Form(""),
    model: str = Form("gpt-5.4-mini"),
    orgId: str = Form(""),
    timeout: str = Form("30"),
    modelParams: str = Form("temperature=0.2")
):
    what = "a notification for the complainant regarding the admissibility outcome" if which == "notify" else "a message for the Investigation Officer regarding the outcome or required next action"
    system_prompt = f"{globalInstructions}\n\nReturn only the requested message text, no quotes, no preamble, not exceeding 200 characters."
    user_text = f"COMPLAINT RECORD\n\n{complaintBlockText}\n\nEXAMINATION\nRecommendation: {recommendation}\nRemarks:\n{remarks}\n\nWrite {what}."

    try:
        out = await call_llm_multimodal(
            endpoint=endpoint,
            api_key=apiKey,
            model=model,
            system_prompt=system_prompt,
            user_text=user_text,
            max_tokens=400,
            org_id=orgId,
            model_params_str=modelParams,
            timeout_sec=int(timeout) if timeout.isdigit() else 30
        )
        out = out.strip().strip('"\'')
        return {"success": True, "text": out}
    except HTTPException as he:
        return JSONResponse(status_code=he.status_code, content={"success": False, "error": he.detail})
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


if __name__ == "__main__":
    print("=" * 65)
    print("  Sindh Ombudsman AI - Scanned PDF & Multimodal Vision Server")
    print("  Serving Application at: http://localhost:8000")
    print("  Supported Formats: Scanned/Digital PDF, DOCX, TXT, JPG, PNG")
    print("=" * 65)
    uvicorn.run(app, host="0.0.0.0", port=8000)
